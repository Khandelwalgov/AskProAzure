import docx
import fitz


DOCX_MIMETYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def clean_text(text):
    return "\n".join(line.rstrip() for line in (text or "").splitlines()).strip()


def detect_heading(text, max_length=120):
    for raw_line in (text or "").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if len(line) <= max_length:
            return line
        return line[:max_length].rstrip() + "..."
    return ""


def make_part(text, metadata):
    cleaned = clean_text(text)
    if not cleaned:
        return None
    part_metadata = dict(metadata)
    part_metadata["section_heading"] = part_metadata.get("section_heading") or detect_heading(cleaned)
    return {"text": cleaned, "metadata": part_metadata}


def extract_pdf_parts(path):
    try:
        parts = []
        with fitz.open(path) as doc:
            page_count = len(doc)
            for page_index, page in enumerate(doc, start=1):
                part = make_part(
                    page.get_text("text"),
                    {
                        "source_type": "pdf",
                        "page": page_index,
                        "page_label": str(page_index),
                        "page_count": page_count,
                        "part_index": page_index - 1,
                    },
                )
                if part:
                    parts.append(part)
        return parts
    except Exception as exc:
        raise RuntimeError(f"Error reading PDF: {exc}") from exc


def extract_docx_parts(path):
    try:
        document = docx.Document(path)
        parts = []
        current_lines = []
        current_heading = ""
        section_index = 0
        paragraph_start = None

        def flush(paragraph_end):
            nonlocal current_lines, current_heading, section_index, paragraph_start
            if not current_lines:
                return
            part = make_part(
                "\n".join(current_lines),
                {
                    "source_type": "docx",
                    "section_heading": current_heading,
                    "section_index": section_index,
                    "paragraph_start": paragraph_start,
                    "paragraph_end": paragraph_end,
                    "part_index": section_index,
                },
            )
            if part:
                parts.append(part)
                section_index += 1
            current_lines = []
            paragraph_start = None

        for index, paragraph in enumerate(document.paragraphs, start=1):
            text = paragraph.text.strip()
            if not text:
                continue

            style_name = (paragraph.style.name if paragraph.style else "").lower()
            is_heading = style_name.startswith("heading")
            if is_heading and current_lines:
                flush(index - 1)

            if paragraph_start is None:
                paragraph_start = index
            if is_heading:
                current_heading = text
            current_lines.append(text)

        flush(len(document.paragraphs))
        return parts
    except Exception as exc:
        raise RuntimeError(f"Error reading DOCX: {exc}") from exc


def extract_txt_parts(path):
    try:
        with open(path, "r", encoding="utf-8") as handle:
            lines = handle.readlines()
    except UnicodeDecodeError:
        with open(path, "r", encoding="utf-8", errors="replace") as handle:
            lines = handle.readlines()
    except Exception as exc:
        raise RuntimeError(f"Error reading TXT: {exc}") from exc

    parts = []
    block = []
    line_start = None
    part_index = 0

    def flush(line_end):
        nonlocal block, line_start, part_index
        if not block:
            return
        part = make_part(
            "".join(block),
            {
                "source_type": "txt",
                "line_start": line_start,
                "line_end": line_end,
                "part_index": part_index,
            },
        )
        if part:
            parts.append(part)
            part_index += 1
        block = []
        line_start = None

    for line_number, line in enumerate(lines, start=1):
        if line.strip():
            if line_start is None:
                line_start = line_number
            block.append(line)
        else:
            flush(line_number - 1)

    flush(len(lines))
    return parts


def extract_document_parts(path, mimetype):
    if mimetype == "application/pdf":
        return extract_pdf_parts(path)
    if mimetype == DOCX_MIMETYPE:
        return extract_docx_parts(path)
    if mimetype == "text/plain":
        return extract_txt_parts(path)

    raise ValueError(f"Unsupported MIME type: {mimetype}")


def extract_text(path, mimetype):
    return "\n\n".join(part["text"] for part in extract_document_parts(path, mimetype)).strip()
